"""
Document Intelligence Service for Deep Research Application

Processes uploaded documents (PDF, DOCX, TXT) and extracts content as markdown.
Uses Azure Document Intelligence for PDF and DOCX, simple text extraction for TXT.
"""

import aiofiles
import structlog
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime, timedelta
from urllib.parse import urlparse
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import DocumentContentFormat
from azure.core.credentials import AzureKeyCredential
from azure.storage.blob.aio import BlobServiceClient
from azure.storage.blob import BlobSasPermissions, generate_blob_sas
from azure.identity.aio import ClientSecretCredential
import aiohttp
from docx import Document as DocxDocument
import tempfile

from ..models.file_models import FileType, FileMetadata

logger = structlog.get_logger(__name__)


class DocumentIntelligenceService:
    """
    Service for extracting content from documents using Azure Document Intelligence.
    
    Supports:
    - PDF: Azure Document Intelligence with markdown output
    - DOCX: python-docx for text extraction
    - TXT: Direct text reading
    """
    
    def __init__(
        self,
        endpoint: str,
        key: str,
        azure_tenant_id: str = None,
        azure_client_id: str = None,
        azure_client_secret: str = None,
        azure_blob_storage_name: str = None
    ):
        """Initialize Document Intelligence service."""
        self.endpoint = endpoint
        self.key = key
        
        # Store Azure credentials for SAS token generation
        self.azure_tenant_id = azure_tenant_id
        self.azure_client_id = azure_client_id
        self.azure_client_secret = azure_client_secret
        self.azure_blob_storage_name = azure_blob_storage_name
        
        # Create Document Intelligence client
        self.client = DocumentIntelligenceClient(
            endpoint=endpoint,
            credential=AzureKeyCredential(key)
        )
        
        logger.info("Document Intelligence service initialized")
    
    async def process_document(
        self,
        file_metadata: FileMetadata
    ) -> Dict[str, Any]:
        """
        Process a document and extract content as markdown.
        
        Args:
            file_metadata: File metadata containing file type and blob URL
            
        Returns:
            Dictionary with markdown_content and metadata
        """
        logger.info(
            "Processing document",
            file_id=file_metadata.id,
            file_type=file_metadata.file_type,
            filename=file_metadata.filename
        )
        
        try:
            if file_metadata.file_type == FileType.PDF:
                return await self._process_pdf(file_metadata)
            elif file_metadata.file_type == FileType.DOCX:
                return await self._process_docx(file_metadata)
            elif file_metadata.file_type == FileType.TXT:
                return await self._process_txt(file_metadata)
            else:
                raise ValueError(f"Unsupported file type: {file_metadata.file_type}")
                
        except Exception as e:
            logger.error(
                "Failed to process document",
                error=str(e),
                file_id=file_metadata.id
            )
            raise
    
    async def _process_pdf(self, file_metadata: FileMetadata) -> Dict[str, Any]:
        """
        Process PDF using Azure Document Intelligence.
        
        Args:
            file_metadata: File metadata
            
        Returns:
            Dictionary with markdown content and metadata
        """
        logger.info("Processing PDF with Document Intelligence", file_id=file_metadata.id)
        
        try:
            # Check if file_path is a blob URL that needs a SAS token
            if file_metadata.file_path.startswith("https://") and self.azure_blob_storage_name:
                # Generate SAS token for the blob
                blob_url_with_sas = await self._generate_sas_url(file_metadata.file_path)
                
                # Create analyze request with SAS-enabled URL
                analyze_request = {
                    "urlSource": blob_url_with_sas
                }
                
                logger.info("Analyzing PDF from blob URL with SAS token", file_id=file_metadata.id)
            else:
                # Use URL directly (for non-blob URLs or if SAS not configured)
                analyze_request = {
                    "urlSource": file_metadata.file_path
                }
                
                logger.info("Analyzing PDF from URL", file_id=file_metadata.id)
            
            # Use Document Intelligence with markdown output format
            poller = self.client.begin_analyze_document(
                model_id="prebuilt-layout",
                body=analyze_request,
                output_content_format=DocumentContentFormat.MARKDOWN
            )
            
            result = poller.result()
            
            # Extract markdown content
            markdown_content = result.content or ""
            
            # Extract metadata
            metadata = {
                "page_count": len(result.pages) if result.pages else 0,
                "word_count": len(markdown_content.split()),
                "has_tables": bool(result.tables),
                "table_count": len(result.tables) if result.tables else 0,
                "has_figures": bool(result.figures),
                "figure_count": len(result.figures) if result.figures else 0
            }
            
            logger.info(
                "PDF processing completed",
                file_id=file_metadata.id,
                page_count=metadata["page_count"],
                word_count=metadata["word_count"]
            )
            
            return {
                "markdown_content": markdown_content,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error("PDF processing failed", error=str(e), file_id=file_metadata.id)
            raise
    
    async def _generate_sas_url(self, blob_url: str) -> str:
        """
        Generate a SAS token for a blob URL.
        
        Args:
            blob_url: The blob URL
            
        Returns:
            Blob URL with SAS token appended
        """
        try:
            # Parse the blob URL
            parsed_url = urlparse(blob_url)
            path_parts = parsed_url.path.lstrip('/').split('/', 1)
            container_name = path_parts[0]
            blob_name = path_parts[1] if len(path_parts) > 1 else ""
            
            # Create credential
            credential = ClientSecretCredential(
                tenant_id=self.azure_tenant_id,
                client_id=self.azure_client_id,
                client_secret=self.azure_client_secret
            )
            
            # Create blob service client
            storage_account_url = f"https://{self.azure_blob_storage_name}.blob.core.windows.net"
            
            async with BlobServiceClient(
                account_url=storage_account_url,
                credential=credential
            ) as blob_service_client:
                # Get user delegation key
                start_time = datetime.utcnow()
                expiry_time = start_time + timedelta(hours=1)
                
                user_delegation_key = await blob_service_client.get_user_delegation_key(
                    key_start_time=start_time,
                    key_expiry_time=expiry_time
                )
                
                # Generate SAS token
                sas_token = generate_blob_sas(
                    account_name=self.azure_blob_storage_name,
                    container_name=container_name,
                    blob_name=blob_name,
                    user_delegation_key=user_delegation_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=expiry_time
                )
                
                blob_url_with_sas = f"{blob_url}?{sas_token}"
                logger.info("Generated SAS token for blob", blob_name=blob_name)
                return blob_url_with_sas
                
        except Exception as e:
            logger.error("Failed to generate SAS token", error=str(e))
            raise
    
    async def _process_docx(self, file_metadata: FileMetadata) -> Dict[str, Any]:
        """
        Process DOCX file using python-docx.
        
        Args:
            file_metadata: File metadata
            
        Returns:
            Dictionary with markdown content and metadata
        """
        logger.info("Processing DOCX file", file_id=file_metadata.id)
        
        try:
            # Download file from blob storage
            async with aiohttp.ClientSession() as session:
                async with session.get(file_metadata.file_path) as response:
                    if response.status != 200:
                        raise Exception(f"Failed to download file: {response.status}")
                    
                    content = await response.read()
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                tmp_file.write(content)
                tmp_path = tmp_file.name
            
            try:
                # Extract text using python-docx
                doc = DocxDocument(tmp_path)
                
                # Extract paragraphs and build markdown
                markdown_lines = []
                word_count = 0
                
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        # Simple markdown formatting based on style
                        if paragraph.style.name.startswith('Heading'):
                            # Convert heading styles to markdown headers
                            level = 1
                            if 'Heading' in paragraph.style.name:
                                try:
                                    level = int(paragraph.style.name.split()[-1])
                                except:
                                    level = 1
                            markdown_lines.append(f"{'#' * level} {text}\n")
                        else:
                            markdown_lines.append(f"{text}\n")
                        
                        word_count += len(text.split())
                
                # Extract tables
                table_count = len(doc.tables)
                if table_count > 0:
                    markdown_lines.append("\n## Tables\n")
                    for i, table in enumerate(doc.tables):
                        markdown_lines.append(f"\n### Table {i + 1}\n")
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            markdown_lines.append("| " + " | ".join(cells) + " |\n")
                
                markdown_content = "\n".join(markdown_lines)
                
                # Metadata
                metadata = {
                    "paragraph_count": len(doc.paragraphs),
                    "word_count": word_count,
                    "table_count": table_count,
                    "section_count": len([p for p in doc.paragraphs if p.style.name.startswith('Heading')])
                }
                
                logger.info(
                    "DOCX processing completed",
                    file_id=file_metadata.id,
                    word_count=word_count,
                    table_count=table_count
                )
                
                return {
                    "markdown_content": markdown_content,
                    "metadata": metadata
                }
                
            finally:
                # Clean up temp file
                Path(tmp_path).unlink(missing_ok=True)
                
        except Exception as e:
            logger.error("DOCX processing failed", error=str(e), file_id=file_metadata.id)
            raise
    
    async def _process_txt(self, file_metadata: FileMetadata) -> Dict[str, Any]:
        """
        Process plain text file.
        
        Args:
            file_metadata: File metadata
            
        Returns:
            Dictionary with markdown content and metadata
        """
        logger.info("Processing text file", file_id=file_metadata.id)
        
        try:
            # Download file from blob storage
            async with aiohttp.ClientSession() as session:
                async with session.get(file_metadata.file_path) as response:
                    if response.status != 200:
                        raise Exception(f"Failed to download file: {response.status}")
                    
                    content = await response.text()
            
            # For text files, content is already plain text
            # Just wrap it in markdown code blocks for better formatting
            markdown_content = content
            
            # Calculate metadata
            lines = content.split('\n')
            words = content.split()
            
            metadata = {
                "line_count": len(lines),
                "word_count": len(words),
                "char_count": len(content)
            }
            
            logger.info(
                "Text file processing completed",
                file_id=file_metadata.id,
                word_count=len(words)
            )
            
            return {
                "markdown_content": markdown_content,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error("Text processing failed", error=str(e), file_id=file_metadata.id)
            raise
